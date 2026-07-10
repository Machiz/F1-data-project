import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import os

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_borders(cell, color_hex="CBD5E0", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def clean_math_text(text):
    text = text.replace('\\gamma', 'γ')
    text = text.replace('\\lambda', 'λ')
    text = text.replace('\\le', '≤')
    text = text.replace('\\in', '∈')
    text = text.replace('\\mathcal{D}', 'D')
    text = text.replace('\\{', '{')
    text = text.replace('\\}', '}')
    text = text.replace('\\mid', '|')
    text = text.replace('\\text', '')
    text = text.replace('\\_', '_')
    text = text.replace('{', '')
    text = text.replace('}', '')
    text = text.replace('\\Delta', 'Δ')
    text = text.replace('R^2', 'R²')
    text = text.replace('k^2', 'k²')
    return text

def parse_inline_state(paragraph, text):
    i = 0
    n = len(text)
    bold = False
    italic = False
    code = False
    math = False
    
    current_text = ""
    
    def flush_run():
        nonlocal current_text
        if not current_text:
            return
            
        if math:
            cleaned = clean_math_text(current_text)
            tokens = cleaned.split(' ')
            for t_idx, token in enumerate(tokens):
                if '_' in token:
                    parts = token.split('_')
                    for p_idx, part in enumerate(parts):
                        if not part:
                            continue
                        run = paragraph.add_run(part)
                        run.font.name = 'Cambria Math'
                        run.italic = True
                        if p_idx > 0:
                            run.font.subscript = True
                else:
                    run = paragraph.add_run(token)
                    run.font.name = 'Cambria Math'
                    run.italic = True
                
                if t_idx < len(tokens) - 1:
                    paragraph.add_run(' ')
        else:
            run = paragraph.add_run(current_text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if code:
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(64, 64, 64) # Monospace dark grey for code
                
        current_text = ""
    
    while i < n:
        if text[i:i+2] == '**':
            flush_run()
            bold = not bold
            i += 2
        elif text[i] == '*':
            flush_run()
            italic = not italic
            i += 1
        elif text[i] == '`':
            flush_run()
            code = not code
            i += 1
        elif text[i] == '$':
            flush_run()
            math = not math
            i += 1
        elif text[i] == '\\':
            if i + 1 < n:
                if text[i+1] == '_':
                    current_text += '_'
                else:
                    current_text += text[i+1]
                i += 2
            else:
                current_text += '\\'
                i += 1
        else:
            current_text += text[i]
            i += 1
    flush_run()

def add_custom_heading(doc, text, level):
    h = doc.add_heading(level=level)
    parse_inline_state(h, text)
    h_format = h.paragraph_format
    
    if level == 1:
        h_format.space_before = Pt(24)
        h_format.space_after = Pt(12)
        h_format.keep_with_next = True
        for run in h.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(51, 51, 51) # Charcoal
            run.bold = True
    elif level == 2:
        h_format.space_before = Pt(18)
        h_format.space_after = Pt(8)
        h_format.keep_with_next = True
        for run in h.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(51, 51, 51) # Charcoal
            run.bold = True
    elif level == 3:
        h_format.space_before = Pt(12)
        h_format.space_after = Pt(6)
        h_format.keep_with_next = True
        for run in h.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(102, 102, 102) # Medium grey
            run.bold = True
    return h

def add_math_block_1(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    math_xml = '''
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r>
        <m:t>predicted_cost_of_staying = wait_laps x (predicted_future_pace - current_lap_duration)</m:t>
      </m:r>
    </m:oMath>
    '''
    p._p.append(parse_xml(math_xml))

def add_math_block_2(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    math_xml = '''
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r>
        <m:t>Reward = </m:t>
      </m:r>
      <m:f>
        <m:num>
          <m:r><m:t>median_pace - lap_duration</m:t></m:r>
        </m:num>
        <m:den>
          <m:r><m:t>10.0</m:t></m:r>
        </m:den>
      </m:f>
    </m:oMath>
    '''
    p._p.append(parse_xml(math_xml))

def build_docx(md_path, docx_path, figures_dir):
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(51, 51, 51) # Soft black (#333333)
    
    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    # Cover page parsing flag
    in_cover = True
    cover_lines = []
    
    # Content lines list
    content_lines = []
    
    for line in lines:
        if in_cover:
            if line.strip().startswith('## Table of Contents'):
                in_cover = False
                content_lines.append(line)
            else:
                cover_lines.append(line)
        else:
            content_lines.append(line)
            
    # Process Cover Page
    title_text = ""
    subtitle_lines = []
    for line in cover_lines:
        line_str = line.strip()
        if not line_str:
            continue
        if line_str.startswith('# '):
            title_text = line_str[2:]
        else:
            subtitle_lines.append(line_str)
            
    # Add Cover Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(20)
    title_run = title_p.add_run(title_text)
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(51, 51, 51) # Charcoal
    
    # Add Subtitles
    for s_line in subtitle_lines:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(8)
        clean_text = s_line.replace('**', '').replace('*', '')
        run = sub_p.add_run(clean_text)
        run.font.name = 'Calibri'
        if "Peruvian University" in clean_text:
            run.font.size = Pt(14)
            run.bold = True
        elif "School of Computer Science" in clean_text or "NRC:" in clean_text:
            run.font.size = Pt(11)
            run.italic = True
        else:
            run.font.size = Pt(11)
            
    doc.add_page_break()
    
    i = 0
    num_lines = len(content_lines)
    mermaid_count = 0
    
    while i < num_lines:
        line = content_lines[i]
        line_strip = line.strip()
        
        if not line_strip:
            i += 1
            continue
            
        # 1. Ignore Horizontal Rule
        if line_strip == '---':
            i += 1
            continue
            
        # 2. Heading
        if line_strip.startswith('#'):
            m = re.match(r'^(#+)\s*(.*)$', line_strip)
            if m:
                level = len(m.group(1))
                heading_text = m.group(2)
                if level == 1 and heading_text == title_text:
                    i += 1
                    continue
                add_custom_heading(doc, heading_text, level)
            i += 1
            continue
            
        # 3. Block Equation
        if line_strip.startswith('$$'):
            if 'predicted' in line_strip:
                add_math_block_1(doc)
            elif 'Reward' in line_strip:
                add_math_block_2(doc)
            i += 1
            continue
            
        # 4. Standard Image Link
        m_img = re.match(r'^!\[(.*?)\]\((.*?)\)$', line_strip)
        if m_img:
            caption = m_img.group(1)
            img_rel_path = m_img.group(2)
            md_dir = os.path.dirname(md_path)
            abs_img_path = os.path.normpath(os.path.join(md_dir, img_rel_path))
            
            if os.path.exists(abs_img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                p.add_run().add_picture(abs_img_path, width=Inches(5.5))
                
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_after = Pt(12)
                cap_run = cap_p.add_run(f"Figure: {caption}")
                cap_run.font.size = Pt(9.5)
                cap_run.italic = True
            else:
                print(f"Warning: Image not found: {abs_img_path}")
            i += 1
            continue
            
        # 5. Mermaid Block
        if line_strip.startswith('```mermaid'):
            while i < num_lines and not content_lines[i].strip() == '```':
                i += 1
            i += 1
            
            mermaid_count += 1
            img_name = ""
            if mermaid_count == 1:
                img_name = "mermaid_pipeline.png"
            elif mermaid_count == 2:
                img_name = "mermaid_1.png"
            elif mermaid_count == 3:
                img_name = "mermaid_overtakes.png"
            elif mermaid_count == 4:
                img_name = "mermaid_2.png"
                
            if img_name:
                img_path = os.path.join(figures_dir, img_name)
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                    p.add_run().add_picture(img_path, width=Inches(5.5))
                    
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_after = Pt(12)
                    cap_run = cap_p.add_run(f"Figure: Unified Strategic Flow Diagram ({img_name})" if "pipeline" in img_name or "1" in img_name else f"Figure: Network Aggression Overlay / RL Structure ({img_name})")
                    cap_run.font.size = Pt(9.5)
                    cap_run.italic = True
            continue
            
        # 6. Code Block
        if line_strip.startswith('```'):
            code_lines = []
            i += 1
            while i < num_lines and not content_lines[i].strip().startswith('```'):
                code_lines.append(content_lines[i])
                i += 1
            i += 1
            
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            set_cell_shading(cell, "F7FAFC")
            set_cell_borders(cell, color_hex="CBD5E0", sz="4")
            
            cell_p = cell.paragraphs[0]
            cell_p.paragraph_format.space_before = Pt(4)
            cell_p.paragraph_format.space_after = Pt(4)
            cell_p.paragraph_format.line_spacing = 1.1
            
            code_content = "".join(code_lines)
            run = cell_p.add_run(code_content.strip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(74, 85, 104) # Slate
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue
            
        # 7. Blockquote
        if line_strip.startswith('>'):
            quote_text = line_strip[1:].strip()
            i += 1
            while i < num_lines and content_lines[i].strip().startswith('>'):
                quote_text += "\n" + content_lines[i].strip()[1:].strip()
                i += 1
                
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            set_cell_shading(cell, "EDF2F7")
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="none"/>
                    <w:left w:val="single" w:sz="24" w:space="0" w:color="555555"/>
                    <w:bottom w:val="none"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            ''')
            tcPr.append(tcBorders)
            
            cell_p = cell.paragraphs[0]
            cell_p.paragraph_format.space_before = Pt(6)
            cell_p.paragraph_format.space_after = Pt(6)
            
            run = cell_p.add_run(quote_text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(74, 85, 104) # Slate
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue
            
        # 8. Table
        if line_strip.startswith('|'):
            table_lines = []
            while i < num_lines and content_lines[i].strip().startswith('|'):
                table_lines.append(content_lines[i].strip())
                i += 1
                
            if len(table_lines) < 2:
                continue
                
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            col_count = len(headers)
            
            data_rows = []
            for t_line in table_lines[2:]:
                cells = [c.strip() for c in t_line.split('|')[1:-1]]
                if len(cells) < col_count:
                    cells += [""] * (col_count - len(cells))
                else:
                    cells = cells[:col_count]
                data_rows.append(cells)
                
            tbl = doc.add_table(rows=1 + len(data_rows), cols=col_count)
            tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            
            tblPr = tbl._tbl.tblPr
            tblCellMar = parse_xml(f'''
                <w:tblCellMar {nsdecls("w")}>
                    <w:top w:w="120" w:type="dxa"/>
                    <w:bottom w:w="120" w:type="dxa"/>
                    <w:left w:w="150" w:type="dxa"/>
                    <w:right w:w="150" w:type="dxa"/>
                </w:tblCellMar>
            ''')
            tblPr.append(tblCellMar)
            
            hdr_cells = tbl.rows[0].cells
            for col_idx, h_text in enumerate(headers):
                cell = hdr_cells[col_idx]
                set_cell_shading(cell, "333333") # Charcoal background
                set_cell_borders(cell, color_hex="CBD5E0", sz="4")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                
                clean_h = h_text.replace('**', '').replace('__', '')
                run = p.add_run(clean_h)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                
            for row_idx, r_data in enumerate(data_rows):
                row_cells = tbl.rows[row_idx + 1].cells
                bg_color = "F7FAFC" if row_idx % 2 == 1 else "FFFFFF"
                for col_idx, cell_value in enumerate(r_data):
                    cell = row_cells[col_idx]
                    set_cell_shading(cell, bg_color)
                    set_cell_borders(cell, color_hex="E2E8F0", sz="4")
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    
                    parse_inline_state(p, cell_value)
                    for run in p.runs:
                        run.font.size = Pt(9.5)
                        
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(8)
            continue
            
        # 9. Unordered List Item
        if line_strip.startswith('* ') or line_strip.startswith('- '):
            list_text = line_strip[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_state(p, list_text)
            i += 1
            continue
            
        # 10. Ordered List Item
        m_list = re.match(r'^(\d+)\.\s*(.*)$', line_strip)
        if m_list:
            list_text = m_list.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_state(p, list_text)
            i += 1
            continue
            
        # 11. Normal Paragraph
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_after = Pt(6)
        p_format.line_spacing = 1.15
        
        parse_inline_state(p, line_strip)
        i += 1
        
    doc.save(docx_path)
    print(f"Successfully created: {docx_path}")

if __name__ == '__main__':
    md = r'c:\Users\User\Documents\GitHub\F1-data-project\project\reports\final_report.md'
    docx_out = r'c:\Users\User\Documents\GitHub\F1-data-project\project\reports\final_report.docx'
    fig_dir = r'c:\Users\User\Documents\GitHub\F1-data-project\project\reports\figures'
    
    build_docx(md, docx_out, fig_dir)
